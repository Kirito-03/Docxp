import pandas as pd
from docx import Document

# 1. Crear Excel de prueba
print("Creando datos_prueba.xlsx...")
df = pd.DataFrame({
    "nombre": ["Juan Pérez", "Ana Gómez", "Luis Silva"],
    "departamento": ["Ventas", "Marketing", "Soporte"],
    "ventas": [15000, 12000, 9500]
})
df.to_excel("datos_prueba.xlsx", index=False)
print("Excel creado exitosamente.")

# 2. Crear Plantilla Word de prueba (.docx) para docxtpl
print("Creando plantilla_prueba.docx...")
doc = Document()
doc.add_heading('Reporte de Rendimiento', 0)

doc.add_paragraph('Hola {{ nombre }},')
doc.add_paragraph('Este es un reporte automático generado por Docxp.')
doc.add_paragraph('Hemos detectado que en el departamento de {{ departamento }} has alcanzado un total de ${{ ventas }} en ventas este mes.')

doc.add_heading('Análisis de Inteligencia Artificial:', level=2)
# Aquí se inyectará el contenido generado por DeepSeek
doc.add_paragraph('{{ ai_content }}')

doc.add_paragraph('Saludos,\nEl equipo de Docxp.')

doc.save("plantilla_prueba.docx")
print("Plantilla Word creada exitosamente.")
