"""Document generation and preview endpoints."""

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app.api.deps import get_db
from app.models.document import Document
from app.schemas.document import (
    FieldFormat,
    GenerateRequest,
    PreviewRequest,
    DocumentResponse,
    DocumentListResponse,
)
from app.services.ai_service import ai_service
from app.services.document_service import DocumentService
from app.services.excel_service import ExcelService
from app.services.template_service import TemplateService

router = APIRouter(prefix="/documents", tags=["Documents"])


def apply_field_format(value: str, fmt: FieldFormat | None) -> str:
    """Apply formatting rules to a string value before injecting into template."""
    if fmt is None:
        return value
    if fmt.zfill is not None and fmt.zfill > 0:
        # Preserve only the numeric part for zfill, then pad
        value = value.zfill(fmt.zfill)
    return value


@router.post("/preview")
async def preview_document(data: PreviewRequest, db: Session = Depends(get_db)):
    """
    Generate an HTML preview of the document without saving.
    Uses AI to generate content based on Excel data and user prompt.
    """
    import docx
    # Get template info
    template = TemplateService.get_template(db, data.template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")

    # Prepare context from Excel data
    context_text = ""
    if data.excel_data:
        context_text = ExcelService.prepare_context_for_ai(data.excel_data)

    # Generate content with AI
    ai_content = ""
    if data.ai_prompt or context_text:
        try:
            ai_content = await ai_service.generate_content(
                context=context_text,
                user_prompt=data.ai_prompt or "Genera el contenido del documento.",
                category=template.category,
                template_instructions=template.ai_instructions,
            )
        except (ConnectionError, RuntimeError) as e:
            ai_content = f"[Error de IA: {str(e)}]\n\nContenido de ejemplo para previsualización."

    # Build template context
    template_context = {
        "titulo": "Previsualización",
        "contenido": ai_content,
    }

    # Add Excel data to context if available
    if data.excel_data and "data" in data.excel_data and data.excel_data["data"]:
        first_sheet = list(data.excel_data["data"].keys())[0]
        rows = data.excel_data["data"][first_sheet]
        if rows:
            first_row = rows[0]
            template_context["datos"] = rows
            
            # Apply custom field mappings + formatting using the first row
            for var_name, col_name in (data.custom_fields or {}).items():
                raw = str(first_row[col_name]) if col_name in first_row else col_name
                fmt = (data.field_formats or {}).get(var_name)
                template_context[var_name] = apply_field_format(raw, fmt)
        else:
            template_context.update(data.custom_fields or {})
    else:
        template_context.update(data.custom_fields or {})

    # Fix OS path separators for Docker
    safe_path = template.file_path.replace("\\", "/")

    # Generate the document to a buffer
    buffer, _ = DocumentService.generate_document(
        template_path=safe_path,
        context=template_context,
        output_filename="preview.docx",
    )

    # Extract text from the generated document
    doc = docx.Document(buffer)
    extracted_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    # Generate HTML preview
    html = DocumentService.generate_preview_html(
        ai_content=extracted_text,
        template_name=template.name,
        title=f"Preview - {template.name}",
        excel_summary=data.excel_data,
    )

    return {"html": html, "ai_content": ai_content}


@router.post("/generate", response_model=DocumentResponse, status_code=201)
async def generate_document(data: GenerateRequest, db: Session = Depends(get_db)):
    """
    Generate a complete document:
    1. Process Excel data context
    2. Generate content with AI
    3. Render into .docx template
    4. Save to database and filesystem
    """
    # Get template
    template = TemplateService.get_template(db, data.template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")

    # Create document record
    doc = Document(
        template_id=template.id,
        title=data.title,
        excel_data=data.excel_data,
        ai_prompt=data.ai_prompt,
        status="generating",
    )
    db.add(doc)
    db.commit()

    try:
        # Prepare AI context
        context_text = ""
        if data.excel_data:
            context_text = ExcelService.prepare_context_for_ai(data.excel_data)

        # Generate AI content
        ai_content = ""
        if data.ai_prompt or context_text:
            ai_content = await ai_service.generate_content(
                context=context_text,
                user_prompt=data.ai_prompt or "Genera el contenido del documento.",
                category=template.category,
                template_instructions=template.ai_instructions,
            )

        # Build template context
        template_context = {
            "titulo": data.title,
            "contenido": ai_content,
        }

        # Add Excel data to context if available
        if data.excel_data and "data" in data.excel_data and data.excel_data["data"]:
            first_sheet = list(data.excel_data["data"].keys())[0]
            rows = data.excel_data["data"][first_sheet]
            if rows:
                first_row = rows[0]
                template_context["datos"] = rows
                
                # Apply custom field mappings + formatting using the first row
                for var_name, col_name in (data.custom_fields or {}).items():
                    raw = str(first_row[col_name]) if col_name in first_row else col_name
                    fmt = (data.field_formats or {}).get(var_name)
                    template_context[var_name] = apply_field_format(raw, fmt)
            else:
                template_context.update(data.custom_fields or {})
        else:
            template_context.update(data.custom_fields or {})

        # Fix OS path separators for Docker
        safe_path = template.file_path.replace("\\", "/")

        # Generate the document
        buffer, filename = DocumentService.generate_document(
            template_path=safe_path,
            context=template_context,
            output_filename=f"{data.title.replace(' ', '_')}.docx",
        )

        # Save to filesystem
        file_path = DocumentService.save_document(buffer, filename)

        # Generate HTML preview
        html = DocumentService.generate_preview_html(
            ai_content=ai_content,
            template_name=template.name,
            title=data.title,
        )

        # Update document record
        doc.ai_response = ai_content
        doc.content_html = html
        doc.generated_file_path = file_path
        doc.status = "completed"
        db.commit()
        db.refresh(doc)

        return DocumentResponse.model_validate(doc)

    except Exception as e:
        doc.status = "error"
        db.commit()
        raise HTTPException(
            status_code=500,
            detail=f"Error generando documento: {str(e)}",
        )


@router.get("/{document_id}/download")
def download_document(document_id: str, db: Session = Depends(get_db)):
    """Download a generated document as .docx."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    if not doc.generated_file_path:
        raise HTTPException(status_code=404, detail="Archivo no disponible")

    try:
        with open(doc.generated_file_path, "rb") as f:
            content = f.read()

        import io
        buffer = io.BytesIO(content)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{doc.title}.docx"'
            },
        )
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Archivo no encontrado en disco")

from fastapi import Form, File, UploadFile
import json
import io
import zipfile

@router.post("/batch-generate")
async def batch_generate_document(
    template_id: str = Form(...),
    title: str = Form(...),
    custom_fields: str = Form("{}"),
    ai_prompt: str = Form(None),
    selected_rows: str = Form(None),
    sheet_name: str = Form(None),
    field_formats: str = Form("{}"),
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """
    Generate a ZIP file with multiple documents, one for each Excel row.
    Accepts optional sheet_name to restrict processing to a single sheet.
    """
    template = TemplateService.get_template(db, template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Plantilla no encontrada")

    if not file.filename or not file.filename.endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="Se requiere un archivo Excel válido para la generación masiva")

    from app.services.excel_service import ExcelService
    from jinja2 import Template as JinjaTemplate
    
    # Process the excel file using the robust ExcelService
    try:
        # Seek to beginning in case it was read
        await file.seek(0)
        excel_data = await ExcelService.extract_data(file, sheet_name=sheet_name)
        rows = []
        # If a sheet_name was specified, only iterate that sheet to avoid mixing data
        target_sheets = [sheet_name] if sheet_name and sheet_name in excel_data.get("data", {}) else list(excel_data.get("data", {}).keys())
        for sheet in target_sheets:
            rows.extend(excel_data["data"][sheet])
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error leyendo el Excel: {str(e)}")

    if not rows:
        raise HTTPException(status_code=400, detail="El archivo Excel está vacío o no contiene filas con datos válidos")

    # Filter rows if selected_rows is provided
    if selected_rows:
        try:
            selected_indices = json.loads(selected_rows)
            if isinstance(selected_indices, list):
                rows = [rows[i] for i in selected_indices if i < len(rows)]
        except Exception:
            pass

    if not rows:
        raise HTTPException(status_code=400, detail="No hay filas seleccionadas para procesar")

    # Parse custom fields
    # Parse custom fields and field_formats
    try:
        custom_fields_dict = json.loads(custom_fields)
    except Exception:
        custom_fields_dict = {}

    field_formats_raw: dict = {}
    try:
        field_formats_raw = json.loads(field_formats)
    except Exception:
        field_formats_raw = {}
    # Convert raw dicts to FieldFormat objects
    field_formats_dict = {
        k: FieldFormat(**v) if isinstance(v, dict) else FieldFormat()
        for k, v in field_formats_raw.items()
    }

    zip_buffer = io.BytesIO()
    seen_names = {}
    
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for idx, row in enumerate(rows):
            # Build context. `custom_fields_dict` contains mapping from template variable to Excel column name
            row_context = {}
            for var_name, col_name in custom_fields_dict.items():
                raw = str(row[col_name]) if col_name in row else col_name
                fmt = field_formats_dict.get(var_name)
                row_context[var_name] = apply_field_format(raw, fmt)
                    
            template_context = {
                "titulo": title,
                **row,          # Raw excel columns
                **row_context   # Explicit mappings override raw excel columns
            }
            
            # Use empty string for "contenido" as batch skips AI generation for speed
            template_context["contenido"] = ""

            # Parse template title with row context (in case it has variables like {{Nombre}})
            try:
                parsed_title = JinjaTemplate(title).render(**template_context)
            except Exception:
                parsed_title = title
            
            safe_title = parsed_title.replace(' ', '_')
            base_filename = f"{safe_title}.docx"
            if base_filename in seen_names:
                seen_names[base_filename] += 1
                final_filename = f"{safe_title}_{seen_names[base_filename]}.docx"
            else:
                seen_names[base_filename] = 1
                final_filename = base_filename
            
            # Fix OS path separators for Docker
            safe_path = template.file_path.replace("\\", "/")

            out_buffer, filename = DocumentService.generate_document(
                template_path=safe_path,
                context=template_context,
                output_filename=final_filename,
            )
            
            zip_file.writestr(filename, out_buffer.getvalue())
            
    zip_buffer.seek(0)
    
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={
            "Content-Disposition": f'attachment; filename="{title.replace(" ", "_")}_batch.zip"'
        }
    )


@router.get("/", response_model=DocumentListResponse)
def list_documents(
    skip: int = 0,
    limit: int = 50,
    status: str | None = None,
    db: Session = Depends(get_db),
):
    """List all generated documents."""
    query = db.query(Document)
    if status:
        query = query.filter(Document.status == status)

    total = query.count()
    items = (
        query.order_by(Document.created_at.desc())
        .offset(skip)
        .limit(limit)
        .all()
    )
    return DocumentListResponse(
        items=[DocumentResponse.model_validate(d) for d in items],
        total=total,
    )
