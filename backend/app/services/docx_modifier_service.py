import json
import logging
from docx import Document

logger = logging.getLogger("docxp")

class DocxModifierService:
    @staticmethod
    def apply_replacements(file_path: str, rules_json: str):
        """
        Applies replacement rules to a DOCX file while preserving formatting.
        rules_json expects: '[{"search": "text", "replace": "{{ var }}"}]'
        """
        try:
            rules = json.loads(rules_json)
            if not rules:
                return
        except Exception:
            logger.error("Failed to parse replacement rules JSON.")
            return

        try:
            doc = Document(file_path)
            
            for rule in rules:
                search_str = rule.get("search")
                replace_str = rule.get("replace")
                if not search_str or not replace_str:
                    continue
                
                # Replace in standard paragraphs
                for p in doc.paragraphs:
                    DocxModifierService._replace_in_paragraph(p, search_str, replace_str)
                
                # Replace in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                DocxModifierService._replace_in_paragraph(p, search_str, replace_str)
            
            doc.save(file_path)
            logger.info(f"Applied {len(rules)} replacement rules to {file_path}")
            
        except Exception as e:
            logger.error(f"Error applying replacements to {file_path}: {e}")
            raise

    @staticmethod
    def _replace_in_paragraph(paragraph, search_str: str, replace_str: str):
        """
        Replaces text within a paragraph, preserving runs if possible.
        """
        if search_str not in paragraph.text:
            return

        # 1. Attempt to replace perfectly within individual runs
        replaced_in_run = False
        for run in paragraph.runs:
            if search_str in run.text:
                run.text = run.text.replace(search_str, replace_str)
                replaced_in_run = True
                
        if replaced_in_run:
            return

        # 2. Fallback: Word spans multiple runs (e.g. spelling checks split runs)
        # We replace the text in the first run and clear the rest.
        # This loses inline formatting but guarantees the text is replaced.
        if search_str in paragraph.text:
            new_text = paragraph.text.replace(search_str, replace_str)
            if paragraph.runs:
                # Retain the style of the first run
                first_run_style = paragraph.runs[0].style
                for i in range(len(paragraph.runs)):
                    paragraph.runs[i].text = ""
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)
